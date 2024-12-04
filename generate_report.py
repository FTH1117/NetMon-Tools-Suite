import os
import re
import argparse
import ticket_fetcher
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

def load_customer_details(customer_dir):
    details_path = os.path.join(customer_dir, 'customer_details.txt')
    details = {}
    print(f"Loading customer details from: {details_path}")
    with open(details_path, 'r') as f:
        for line in f:
            key, value = line.strip().split(': ', 1)
            details[key.strip()] = value.strip()
    print(f"Loaded customer details: {details}")
    return details

def find_section_paragraph(doc, section_title):
    for i, paragraph in enumerate(doc.paragraphs):
        if section_title in paragraph.text:
            print(f"Found section '{section_title}' at paragraph index: {i}")
            return i
    print(f"Section '{section_title}' not found.")
    return None

def collect_host_sla_data(month_dir):
    sla_data = {}
    for host_dir_name in sorted(os.listdir(month_dir)):
        host_dir = os.path.join(month_dir, host_dir_name)
        if os.path.isdir(host_dir):
            # Look for the SLA text file
            for file_name in os.listdir(host_dir):
                if file_name.endswith('.txt') and 'SLA' in file_name:
                    sla_file_path = os.path.join(host_dir, file_name)
                    with open(sla_file_path, 'r') as f:
                        sla_content = f.read()
                    sla_percentage_match = re.search(r"Uptime: ([\d\.]+)%", sla_content)
                    if sla_percentage_match:
                        sla_percentage = sla_percentage_match.group(1)
                        sla_data[host_dir_name] = sla_percentage
                    else:
                        print(f"Uptime percentage not found in {sla_file_path}")
    return sla_data


def insert_host_and_graphs(doc, host_dir, host_full_name, insert_index, sla_data, insert_host_name=True, insert_graph_name=True):
    # Check if there are any PNG files in the directory
    graph_files = [f for f in os.listdir(host_dir) if f.endswith('.png')]
    if not graph_files:
        print(f"No graphs found for host '{host_full_name}'. Skipping this host.")
        return insert_index  # Return the original index without modification

    if insert_host_name:
        print(f"Inserting host '{host_full_name}' at index {insert_index}")
        # Insert the host name as a heading before insert_index
        doc.paragraphs[insert_index].insert_paragraph_before(host_full_name, style="Heading 2")
        insert_index += 1  # Move to the next position

        # Get SLA uptime for the host if available
        if sla_data:
            sla_percentage = sla_data.get(host_full_name, None)
            if sla_percentage:
                # Insert SLA uptime after host name
                doc.paragraphs[insert_index].insert_paragraph_before(f"Uptime: {sla_percentage}%")
                insert_index += 1  # Move to the next position

    # Define the keyword order
    keyword_order = {
        "ping": 1,
        "uptime": 2,
        "cpu": 3,
        "memory": 4,
        "disk": 5
    }

    # Sort files based on keyword priorities; unmatched files are sorted normally
    sorted_graphs = sorted(
        graph_files,
        key=lambda x: (keyword_order.get(next((kw for kw in keyword_order if kw in x.lower()), ""), float('inf')), x)
    )

    # Insert each graph in the sorted order
    for filename in sorted_graphs:
        graph_name = re.sub(r"_\d+\.png$", "", filename).replace("_", " ").replace("^", "/")
        print(f"Inserting graph '{graph_name}' from file '{filename}'")

        if insert_graph_name:
            # Insert graph name before insert_index
            doc.paragraphs[insert_index].insert_paragraph_before(graph_name)
            insert_index += 1  # Move to the next position


        # Insert image before insert_index
        new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
        run = new_paragraph.add_run()
        image_path = os.path.join(host_dir, filename)
        run.add_picture(image_path, width=Inches(5))
        print(f"Inserted image '{filename}'")
        insert_index += 1  # Move to the next position

    return insert_index  # Return the updated insert_index



def generate_report(month, year, customer_id):
    base_dir = f"/home/almalinux/{customer_id}"
    month_dir = os.path.join(base_dir, f"{year}-{str(month).zfill(2)}")

    # Load customer details to get Project ID and Project Name
    customer_details = load_customer_details(base_dir)
    project_id = customer_details.get("Project ID", "")
    project_name = customer_details.get("Project Name", "")

    # Format the output file name as per your requirement
    output_filename = f"{project_id} {project_name} - {datetime(year, month, 1).strftime('%B')} {year} Monthly Report.docx"
    output_path = os.path.join(month_dir, output_filename)

    # Determine whether to use the network template or the standard one
    network_dir = os.path.join(month_dir, 'network')
    use_network_template = False
    if os.path.isdir(network_dir):
        # Check if the network directory contains any subdirectories (host directories)
        host_dir_names = [d for d in os.listdir(network_dir) if os.path.isdir(os.path.join(network_dir, d))]
        if host_dir_names:
            use_network_template = True

    if use_network_template:
        template_path = "template.docx"
        print("Using template with network section: template.docx")
    else:
        template_path = "template_xnetwork.docx"
        print("Using template without network section: template_xnetwork.docx")

    print(f"Loading template from: {template_path}")
    doc = Document(template_path)

    # Define replacements
    replacements = {
        "<Project ID>": project_id,
        "<Project Name>": project_name,
        "<Month>": datetime(year, month, 1).strftime("%B"),
        "<Year>": str(year),
        "<Today Date>": datetime.now().strftime("%Y-%m-%d"),
    }

    # Apply replacements in the template with custom styles
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                print(f"Replacing '{key}' with '{value}'")
                for run in paragraph.runs:
                    if key in run.text:
                        # Replace the placeholder with the actual value
                        run.text = run.text.replace(key, value)
                        # Apply custom styles
                        run.font.size = Pt(16)  # Set font size (adjust as needed)
                        run.font.bold = True    # Set text to bold


    # Collect SLA data
    sla_data = collect_host_sla_data(month_dir)

    # Find the paragraph containing the specified line
    specified_line = "reinforcing IP ServerOne's commitment to operational reliability and service availability."
    insert_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if specified_line in paragraph.text:
            print(f"Found specified line at paragraph index: {i}")
            insert_index = i + 1  # Insert after this paragraph
            break

    # Insert host SLA data after the specified line
    if insert_index is not None:
        for host_name, sla_percentage in sla_data.items():
            # Create a new paragraph
            paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            # Add host name run with bold formatting
            run_host = paragraph.add_run(host_name)
            run_host.font.bold = True
            # Add the rest of the text
            run_sla = paragraph.add_run(f": Uptime {sla_percentage}%")
            # Increment insert_index
            insert_index += 1  # Move to the next position
    else:
        print(f"Specified line not found in the document. Cannot insert uptime data.")


    # Insert graphs and host information
    section_index = find_section_paragraph(
        doc,
        "The purpose of this section is to illustrate the performance and capacity trends for each system, supporting effective resource planning and ensuring alignment with operational requirements."
    )
    if section_index is not None:
        insert_index = section_index + 1
        host_dir_names = sorted(os.listdir(month_dir))
        for idx, host_dir_name in enumerate(host_dir_names):
            host_dir = os.path.join(month_dir, host_dir_name)
            if os.path.isdir(host_dir) and host_dir_name != 'network':
                host_full_name = host_dir_name  # Use the full directory name as the host name
                insert_index_before = insert_index
                insert_index = insert_host_and_graphs(doc, host_dir, host_full_name, insert_index, sla_data)
                # Check if content was inserted
            if insert_index > insert_index_before:
                # Insert a page break after each host except the last one
                if idx < len(host_dir_names) - 1:
                    paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                    run = paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    insert_index += 1  # Move to the next position
            else:
                print(f"Host '{host_full_name}' was skipped, not inserting page break.")
    else:
        print("Section for inserting regular graphs not found.")

    # Insert network graphs after specified sentence, if the network template is used
    if use_network_template:
        network_section_sentence = "This analysis helps identify network performance patterns, ensuring optimal resource allocation and supporting proactive capacity planning."
        network_section_index = find_section_paragraph(doc, network_section_sentence)

        if network_section_index is not None:
            insert_index = network_section_index + 1
            if os.path.isdir(network_dir):
                host_dir_names = [d for d in os.listdir(network_dir) if os.path.isdir(os.path.join(network_dir, d))]
                if host_dir_names:
                    for idx, host_dir_name in enumerate(host_dir_names):
                        host_dir = os.path.join(network_dir, host_dir_name)
                        host_full_name = host_dir_name
                        insert_index_before = insert_index
                        insert_index = insert_host_and_graphs(
                            doc, host_dir, host_full_name, insert_index, sla_data=None, insert_host_name=False, insert_graph_name=False
                        )
                        # Check if content was inserted
                        if insert_index > insert_index_before:
                            # Insert a page break after each host except the last one
                            if idx < len(host_dir_names) - 1:
                                paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                                run = paragraph.add_run()
                                run.add_break(WD_BREAK.PAGE)
                                insert_index += 1  # Move to the next position
                        else:
                            print(f"Network host '{host_full_name}' was skipped, not inserting page break.")
                else:
                    print(f"No network hosts found in {network_dir}")
            else:
                print(f"No network directory found at {network_dir}")
        else:
            print("Section for inserting network graphs not found.")
    else:
        print("Skipping network graphs insertion as no network hosts are found.")

    # Insert tickets into the document
    ticket_fetcher.fetch_and_insert_tickets(doc, base_dir, month, year)

    print(f"Saving output to: {output_path}")
    doc.save(output_path)
    print("Report generated successfully.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Managed Service Report")
    parser.add_argument("--month", type=int, required=True, help="Report month (1-12)")
    parser.add_argument("--year", type=int, required=True, help="Report year (e.g., 2024)")
    parser.add_argument("--customer", type=str, required=True, help="Customer Project ID")

    args = parser.parse_args()
    generate_report(args.month, args.year, args.customer)

