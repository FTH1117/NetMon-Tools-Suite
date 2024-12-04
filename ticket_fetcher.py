# ticket_fetcher.py

import requests
from docx import Document
from datetime import datetime, timedelta, timezone
import time
import os
import re
import csv
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement

def fetch_and_insert_tickets(document, customer_dir, month, year):
    # Paths to files
    customer_details_path = os.path.join(customer_dir, 'customer_details.txt')
    
    # SupportPal API credentials
    api_base_url = '<SUPPORTPAL_URL>'
    api_token = '<SUPPORTPAL_API>'  # Replace with your actual API token
    
    # Read Project ID and Subscription IDs from customer_details.txt
    service_id_values = set()
    project_id_num = None
    
    with open(customer_details_path, 'r') as file:
        for line in file:
            line = line.strip()
            if line.startswith('Project ID: '):
                project_id_value = line.split('Project ID: ')[1].strip()
                # Extract numerical part after 'AA0' or 'AA' or 'SS0' or 'SS'
                match = re.search(r'(?:AA0*|AA|SS0*|SS)(\d+)', project_id_value)
                if match:
                    project_id_num = match.group(1).lstrip('0')
                else:
                    print(f"Could not extract numerical Service ID from Project ID: {project_id_value}")
                    project_id_num = None
            elif line.startswith('Subscription ID'):
                # Extract value after ':'
                _, value = line.split(':', 1)
                subscription_id_value = value.strip()
                # Extract numerical part
                match = re.search(r'(?:AA0*|AA|SS0*|SS)(\d+)', subscription_id_value)
                if match:
                    subscription_id_num = match.group(1).lstrip('0')
                    service_id_values.add(subscription_id_num)
                else:
                    print(f"Could not extract numerical Service ID from Subscription ID: {subscription_id_value}")
    
    if not service_id_values:
        if project_id_num:
            service_id_values.add(project_id_num)
        else:
            print("No 'Subscription ID' or 'Project ID' found in customer_details.txt.")
            return
    
    # Remove any service ID values that are the same as the Project ID
    if project_id_num in service_id_values:
        service_id_values.remove(project_id_num)
    
    # Now, add the project_id_num
    if project_id_num:
        service_id_values.add(project_id_num)
    
    # Convert service_id_values to a list
    service_id_values = list(service_id_values)
    
    # Define the UTC+8 timezone
    utc_plus_8 = timezone(timedelta(hours=8))
    
    # Compute start and end timestamps for the given month and year
    try:
        # Start date is the first day of the month at 00:00:00 UTC+8
        start_date = datetime(year, month, 1, tzinfo=utc_plus_8)
        # End date is the last day of the month at 23:59:59 UTC+8
        if month == 12:
            end_date = datetime(year + 1, 1, tzinfo=utc_plus_8) - timedelta(seconds=1)
        else:
            end_date = datetime(year, month + 1, 1, tzinfo=utc_plus_8) - timedelta(seconds=1)
    
    except Exception as e:
        print(f"Error computing date range: {e}")
        return
    
    # Initialize 'limit'
    limit = 100
    
    # The custom field ID for "Service ID"
    CUSTOM_FIELD_ID = 6  # Use the actual custom field ID
    
    try:
        # Set up Basic Authentication with API token as username and 'X' as password
        auth = (api_token, 'X')
    
        # Fetch tickets and implement pagination
        tickets = []
        ticket_ids = set()  # To avoid duplicates
    
        for service_id_value in service_id_values:
            start = 1
            while True:
                # Prepare API request parameters for this service ID
                params = {
                    'order_column': 'created_at',
                    'order_direction': 'asc',
                    'limit': limit,
                    'internal': 0,  # Fetch non-internal tickets
                    'with': 'user',  # Include user data
                    'start': start,
                }
                params[f'customfield[{CUSTOM_FIELD_ID}]'] = service_id_value
    
                response = requests.get(f'{api_base_url}/ticket/ticket', params=params, auth=auth)
    
                # Print the request URL for debugging
                print(f"Request URL: {response.url}")
    
                # Raise exception for HTTP errors
                response.raise_for_status()
                tickets_data = response.json()
    
                # Check for API errors
                if tickets_data.get('status') == 'error':
                    print(f"API Error: {tickets_data.get('message')}")
                    return
    
                page_tickets = tickets_data.get('data', [])
                if not page_tickets:
                    break
    
                for ticket in page_tickets:
                    ticket_id = ticket.get('id')
                    if ticket_id not in ticket_ids:
                        tickets.append(ticket)
                        ticket_ids.add(ticket_id)
    
                # Increment 'start' by the number of tickets received
                start += len(page_tickets)
    
                # If the number of tickets returned is less than 'limit', we've reached the last page
                if len(page_tickets) < limit:
                    break
    
                # Sleep briefly to avoid hitting rate limits
                time.sleep(0.5)
    
        # Function to parse 'created_at' into datetime object in UTC+8
        def parse_created_at(created_at):
            try:
                if isinstance(created_at, int):
                    # Normalize 'created_at' to seconds if necessary
                    if created_at > 1e12:  # Likely in milliseconds
                        created_at_sec = created_at // 1000
                    else:
                        created_at_sec = created_at
                    # Convert timestamp to datetime in UTC, then convert to UTC+8
                    dt_utc = datetime.utcfromtimestamp(created_at_sec).replace(tzinfo=timezone.utc)
                    dt_utc_plus_8 = dt_utc.astimezone(utc_plus_8)
                    return dt_utc_plus_8
                else:
                    # Handle other formats if necessary
                    return None
            except Exception as e:
                print(f"Error parsing 'created_at': {e}")
                return None
    
        # Filter tickets based on creation date in UTC+8
        filtered_tickets = []
        for ticket in tickets:
            created_at = ticket.get('created_at')
            created_at_dt = parse_created_at(created_at)
            if created_at_dt is not None:
                if start_date <= created_at_dt <= end_date:
                    ticket['created_at_dt'] = created_at_dt  # Store the datetime object in the ticket
                    filtered_tickets.append(ticket)
            else:
                # Decide what to do if 'created_at' is None or couldn't be parsed
                pass
        
        if not filtered_tickets:
            print(f"No tickets found for Service IDs: {service_id_values} in {month}/{year}")
            # Instead of returning, proceed to insert the sentence into the document
            # Find the paragraph where we want to insert after
            for i, paragraph in enumerate(document.paragraphs):
                if "help to keep track the progress of issued and requests raised and assess the responsiveness of the support team." in paragraph.text:
                    insert_index = i + 1  # We will insert after this paragraph
                    break
            else:
                print("The phrase was not found in the document.")
                return

            # Get the paragraph where we want to insert after
            p = document.paragraphs[insert_index - 1]

            # Get the document body
            body = document._body
            body_elm = body._body
            p_index = body_elm.index(p._p)

            # Insert a new paragraph with the sentence
            new_paragraph = document.add_paragraph("No ticket has been raised on this month.")
            body_elm.insert(p_index + 1, new_paragraph._p)

            # Insert a page break after the sentence
            page_break_paragraph = document.add_paragraph()
            page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
            body_elm.insert(p_index + 2, page_break_paragraph._p)

            print(f"Inserted message into the document indicating no tickets were raised in {month}/{year}.")
            return  # Exit the function after inserting the message
        else:
            # Sort the filtered tickets by 'created_at_dt'
            filtered_tickets.sort(key=lambda ticket: ticket['created_at_dt'])

            # Write tickets to CSV file
            month_dir = os.path.join(customer_dir, f"{year}-{str(month).zfill(2)}")
            os.makedirs(month_dir, exist_ok=True)
            csv_filename = 'Ticket list.csv'
            csv_path = os.path.join(month_dir, csv_filename)

            with open(csv_path, mode='w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['No', 'Number', 'Date Ticket Created', 'Ticket Subject', 'Email']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

                writer.writeheader()
                for idx, ticket in enumerate(filtered_tickets, start=1):
                    ticket_number = ticket.get('number', 'N/A')
                    ticket_subject = ticket.get('subject', 'N/A')
                    email = ticket.get('user', {}).get('email', 'N/A')
                    created_at_dt = ticket.get('created_at_dt')

                    if created_at_dt:
                        created_at_readable = created_at_dt.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        created_at_readable = 'N/A'

                    writer.writerow({
                        'No': idx,
                        'Number': ticket_number,
                        'Date Ticket Created': created_at_readable,
                        'Ticket Subject': ticket_subject,
                        'Email': email
                    })

            # Now, insert the table into the document
            # Find the paragraph that contains the specific sentence
            for i, paragraph in enumerate(document.paragraphs):
                if "help to keep track the progress of issued and requests raised and assess the responsiveness of the support team." in paragraph.text:
                    insert_index = i + 1  # We will insert after this paragraph
                    break
            else:
                print("The phrase was not found in the document.")
                return

            # Get the paragraph where we want to insert after
            p = document.paragraphs[insert_index - 1]

            # Get the document body
            body = document._body
            body_elm = body._body
            p_index = body_elm.index(p._p)

            # Insert a new paragraph (line break) after the target paragraph
            line_break_paragraph = document.add_paragraph()
            body_elm.insert(p_index + 1, line_break_paragraph._p)

            # Create the table
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Define table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Number'
            hdr_cells[2].text = 'Date Ticket Created'
            hdr_cells[3].text = 'Ticket Subject'
            hdr_cells[4].text = 'Email'  # New column for email

            # Populate the table with ticket data
            for idx, ticket in enumerate(filtered_tickets, start=1):
                row_cells = table.add_row().cells
                ticket_number = ticket.get('number', 'N/A')
                ticket_subject = ticket.get('subject', 'N/A')
                email = ticket.get('user', {}).get('email', 'N/A')  # Extract email
                created_at_dt = ticket.get('created_at_dt')

                if created_at_dt:
                    created_at_readable = created_at_dt.strftime('%Y-%m-%d %H:%M:%S')
                else:
                    created_at_readable = 'N/A'

                row_cells[0].text = str(idx)
                row_cells[1].text = ticket_number
                row_cells[2].text = created_at_readable
                row_cells[3].text = ticket_subject
                row_cells[4].text = email  # Populate email column

            # Insert the table after the line break paragraph
            body_elm.insert(p_index + 2, table._tbl)

            # Insert a page break after the table
            page_break_paragraph = document.add_paragraph()
            page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
            body_elm.insert(p_index + 3, page_break_paragraph._p)

            print(f"Tickets inserted into the document and CSV file '{csv_filename}' created.")
    
    except requests.exceptions.RequestException as e:
        print(f"Error fetching tickets: {e}")
        if e.response:
            print("Response text:", e.response.text)
        return
    
    except Exception as e:
        import traceback
        print(f"An error occurred: {e}")
        traceback.print_exc()
        return
