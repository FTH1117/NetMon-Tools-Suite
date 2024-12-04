import os
import re
import argparse
import ticket_fetcher
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
from docx.enum.text import WD_BREAK
from PIL import Image
import llama_analysis

def load_customer_details(customer_dir):
    details_path = os.path.join(customer_dir, 'customer_details.txt')
    details = {}
    print(f"Loading customer details from: {details_path}")
    with open(details_path, 'r') as f:
        for line in f:
            key, value = line.strip().split(': ', 1)
            details[key.strip()] = value.strip()
    print(f"Loaded customer details: {details}")
    return details

def find_section_paragraph(doc, section_title):
    for i, paragraph in enumerate(doc.paragraphs):
        if section_title in paragraph.text:
            print(f"Found section '{section_title}' at paragraph index: {i}")
            return i
    print(f"Section '{section_title}' not found.")
    return None

def insert_image_with_adjusted_width(run, image_path, original_width_in_inches=6.5, min_width_in_inches=2):
    img = Image.open(image_path)
    height = img.height
    if height > 1800:
        extra_height = height - 1800
        increments = extra_height // 200
        width_reduction = increments * 0.5
        new_width_in_inches = original_width_in_inches - width_reduction
        new_width_in_inches = max(new_width_in_inches, min_width_in_inches)
    else:
        new_width_in_inches = original_width_in_inches
    run.add_picture(image_path, width=Inches(new_width_in_inches))


def generate_grafana_report(month, year, customer_id, llama_selected=False):
    base_dir = f"/home/almalinux/{customer_id}"
    month_dir = os.path.join(base_dir, f"{year}-{str(month).zfill(2)}")

    # Load customer details to get Project ID and Project Name
    customer_details = load_customer_details(base_dir)
    project_id = customer_details.get("Project ID", "")
    project_name = customer_details.get("Project Name", "")

    # Format the output file name
    output_filename = f"{project_id} {project_name} - {datetime(year, month, 1).strftime('%B')} {year} Monthly Report.docx"
    output_path = os.path.join(month_dir, output_filename)

    # Determine whether to use the network template or the standard one
    network_category_dir = None
    for root, dirs, files in os.walk(month_dir):
        if os.path.basename(root).lower() == 'network_traffic':
            network_category_dir = root
            break

    use_network_template = False
    if network_category_dir and os.path.isdir(network_category_dir):
        network_graph_files = [f for f in os.listdir(network_category_dir) if f.endswith('.png')]
        if network_graph_files:
            use_network_template = True

    template_path = "template.docx" if use_network_template else "template_xnetwork.docx"
    print(f"Loading template from: {template_path}")
    doc = Document(template_path)

    # Define replacements
    replacements = {
        "<Project ID>": project_id,
        "<Project Name>": project_name,
        "<Month>": datetime(year, month, 1).strftime("%B"),
        "<Year>": str(year),
        "<Today Date>": datetime.now().strftime("%Y-%m-%d"),
    }

    # Apply replacements in the template with custom styles
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                print(f"Replacing '{key}' with '{value}'")
                for run in paragraph.runs:
                    if key in run.text:
                        # Replace the placeholder with the actual value
                        run.text = run.text.replace(key, value)
                        # Apply custom styles
                        run.font.size = Pt(16)  # Set font size (adjust as needed)
                        run.font.bold = True    # Set text to bold


    # Find the paragraph containing the specified line
    specified_line = "reinforcing IP ServerOne's commitment to operational reliability and service availability."
    insert_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if specified_line in paragraph.text:
            print(f"Found specified line at paragraph index: {i}")
            insert_index = i + 1  # Insert after this paragraph
            break

    if insert_index is not None:
        # Find the Ping Result graph
        ping_result_path = None
        for root, dirs, files in os.walk(month_dir):
            for f in files:
                if f == 'Ping Result.png':
                    ping_result_path = os.path.join(root, f)
                    break
            if ping_result_path:
                break

        if ping_result_path and os.path.isfile(ping_result_path):
            print(f"Inserting Ping Result panel from: {ping_result_path}")
            # Insert picture name (without '.png'), make it bold and bigger
            paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            run = paragraph.add_run("Ping Result")
            run.font.bold = True
            run.font.size = Pt(13)  # Adjust size as needed
            insert_index += 1

            # Insert the picture
            new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            run = new_paragraph.add_run()
            insert_image_with_adjusted_width(run, ping_result_path)
            insert_index += 1
        else:
            print(f"Ping Result graph not found.")

        if llama_selected:
            prompt = llama_analysis.get_system_prompt('Ping_Result')
            print("Performing Llama analysis for Ping Result...")
            analysis_output = llama_analysis.perform_llama_analysis(ping_result_path, prompt)
            if analysis_output:
                # Save the analysis output to a text file in the same directory as the PNG
                analysis_file = os.path.join(os.path.dirname(ping_result_path), "Ping_Result_analysis.txt")
                with open(analysis_file, 'w') as f:
                    f.write(analysis_output)

                # Insert "Overall assessment" in bold
                paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                run = paragraph.add_run("Overall assessment")
                run.font.bold = True
                insert_index += 1

                # Insert the analysis output
                paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                paragraph.text = analysis_output
                insert_index += 1
            else:
                print("Llama analysis failed for Ping Result.")

    else:
        print(f"Specified line not found in the document. Cannot insert Ping Result panel.")

    # Find the section where to insert the graphs
    section_title = "The purpose of this section is to illustrate the performance and capacity trends for each system, supporting effective resource planning and ensuring alignment with operational requirements."
    section_index = find_section_paragraph(doc, section_title)

    if section_index is not None:
        insert_index = section_index + 1

        # Collect graphs from month_dir and its subdirectories, excluding 'Ping Result.png' and 'network' directory
        graph_files = []
        for root, dirs, files in os.walk(month_dir):
            # Skip the 'network' directory if it exists
            if os.path.basename(root).lower() in ['network', 'network_traffic']:
                continue
            for f in files:
                if f.endswith('.png') and f != "Ping Result.png":
                    file_path = os.path.join(root, f)
                    graph_files.append(file_path)

        # Sort graphs based on a predefined order
        keyword_order = {
            "Ping Result": 1,
            "Uptime": 2,
            "CPU utilization": 3,
            "Memory utilization": 4,
            "Disk Space Usage": 5,
            "Network Usage": 6
        }

        def get_sort_order(file_path):
            filename = os.path.basename(file_path)
            graph_name = filename.replace('.png', '')
            return keyword_order.get(graph_name, float('inf'))

        sorted_graphs = sorted(graph_files, key=get_sort_order)


        for idx, file_path in enumerate(sorted_graphs):
            filename = os.path.basename(file_path)
            graph_name = filename.replace('.png', '')
            category = os.path.basename(os.path.dirname(file_path))
            
            # Insert picture name (without '.png'), make it bold and bigger
            paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            run = paragraph.add_run(graph_name)
            run.font.bold = True
            run.font.size = Pt(13)  # Adjust size as needed
            insert_index += 1

            # Insert the picture
            new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            run = new_paragraph.add_run()
            insert_image_with_adjusted_width(run, file_path)
            insert_index += 1

            if llama_selected:
                # Perform Llama analysis
                prompt = llama_analysis.get_system_prompt(category)
                print(f"Performing Llama analysis for {graph_name}...")
                analysis_output = llama_analysis.perform_llama_analysis(file_path, prompt)
                if analysis_output:
                    # Save the analysis output to a text file in the same directory as the PNG
                    analysis_file = os.path.join(os.path.dirname(file_path), f"{graph_name}_analysis.txt")
                    with open(analysis_file, 'w') as f:
                        f.write(analysis_output)

                    # Insert "Overall assessment" in bold
                    paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                    run = paragraph.add_run("Overall assessment")
                    run.font.bold = True
                    insert_index += 1

                    # Insert the analysis output
                    paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                    paragraph.text = analysis_output
                    insert_index += 1
                else:
                    print(f"Llama analysis failed for {graph_name}.")

            # Insert page break after each graph, including the last one
            paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            insert_index += 1

    else:
        print("Section for inserting graphs not found.")

    # Insert network graph if any
    if use_network_template:
        network_section_sentence = "This analysis helps identify network performance patterns, ensuring optimal resource allocation and supporting proactive capacity planning."
        network_section_index = find_section_paragraph(doc, network_section_sentence)
        if network_section_index is not None:
            insert_index = network_section_index + 1
            # Assume network graphs are in network_dir
            network_graph_files = [f for f in os.listdir(network_category_dir) if f.endswith('.png') and os.path.isfile(os.path.join(network_category_dir, f))]
            if network_graph_files:
                # We assume there is only one network graph
                filename = network_graph_files[0]
                # Do not insert graph name for network graph
                # Insert picture
                new_paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                run = new_paragraph.add_run()
                image_path = os.path.join(network_category_dir, filename)
                insert_image_with_adjusted_width(run, image_path)
                insert_index += 1
                
                if llama_selected:
                    # Perform Llama analysis for the network graph
                    category = 'Network_Traffic'  # Or adjust based on your directory naming
                    prompt = llama_analysis.get_system_prompt(category)
                    print("Performing Llama analysis for Network Traffic graph...")
                    analysis_output = llama_analysis.perform_llama_analysis(image_path, prompt)
                    if analysis_output:
                        # Save the analysis output to a text file in the same directory as the PNG
                        analysis_file = os.path.join(network_category_dir, f"{filename.replace('.png', '')}_analysis.txt")
                        with open(analysis_file, 'w') as f:
                            f.write(analysis_output)

                        # Insert "Overall assessment" in bold
                        paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                        run = paragraph.add_run("Overall assessment")
                        run.font.bold = True
                        insert_index += 1

                        # Insert the analysis output
                        paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                        paragraph.text = analysis_output
                        insert_index += 1
                    else:
                        print("Llama analysis failed for Network Traffic graph.")

                # Insert page break after the network graph
                paragraph = doc.paragraphs[insert_index].insert_paragraph_before()
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
                insert_index += 1
            else:
                print("No network graphs found.")
        else:
            print("Network section not found in the template.")
    else:
        print("Skipping network graphs insertion as no network hosts are found.")

    # Insert tickets into the document
    ticket_fetcher.fetch_and_insert_tickets(doc, base_dir, month, year)

    print(f"Saving output to: {output_path}")
    doc.save(output_path)
    print("Grafana report generated successfully.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Grafana Managed Service Report")
    parser.add_argument("--month", type=int, required=True, help="Report month (1-12)")
    parser.add_argument("--year", type=int, required=True, help="Report year (e.g., 2024)")
    parser.add_argument("--customer", type=str, required=True, help="Customer Project ID")
    parser.add_argument("--llama", action='store_true', help="Perform Llama analysis on graphs")

    args = parser.parse_args()
    generate_grafana_report(args.month, args.year, args.customer, args.llama)
